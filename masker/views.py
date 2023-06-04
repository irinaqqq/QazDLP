import io
import os
import docx
from django.conf import settings
from django.http import HttpResponse
from django.shortcuts import render
from django.views.decorators.csrf import csrf_exempt

from .forms import MaskerForm
from .utils import *


def masker(request):
    if request.method == 'POST':
        encryption_type = request.POST.get('encryption_type')
        form = MaskerForm(request.POST, request.FILES)
        if form.is_valid():
            text = form.cleaned_data['text']
            file = request.FILES.get('file', None)

            if encryption_type == 'file':
                if file:
                    _, ext = os.path.splitext(file.name)
                    if ext == '.txt':
                        file_text = file.read().decode('utf-8')
                        file_text = mask_text(file_text)
                        response = HttpResponse(file_text, content_type='text/plain')
                        response['Content-Disposition'] = 'attachment; filename="masked_file.txt"'
                    elif ext == '.docx':
                        doc = docx.Document(file)
                        for para in doc.paragraphs:
                            para.text = mask_text(para.text)
                        output = io.BytesIO()
                        doc.save(output)
                        output.seek(0)
                        response = HttpResponse(output.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                        response['Content-Disposition'] = 'attachment; filename="masked_file.docx"'
                    else:
                        return HttpResponse('Unsupported file type')

                    return response
                else:
                    return HttpResponse('No file was uploaded.')
            elif encryption_type == 'text':
                masked_text = mask_text(text)
                return render(request, 'masker.html', {'form': form, 'masked_text': masked_text, 'text': text})
            elif encryption_type == 'decrypt_file':
                if file:
                    _, ext = os.path.splitext(file.name)
                    if ext == '.txt':
                        file_text = file.read().decode('utf-8')
                        file_text = unmask_text(file_text)
                        response = HttpResponse(file_text, content_type='text/plain')
                        response['Content-Disposition'] = 'attachment; filename="unmasked_file.txt"'
                    elif ext == '.docx':
                        doc = docx.Document(file)
                        for para in doc.paragraphs:
                            para.text = unmask_text(para.text)
                        output = io.BytesIO()
                        doc.save(output)
                        output.seek(0)
                        response = HttpResponse(output.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                        response['Content-Disposition'] = 'attachment; filename="unmasked_file.docx"'
                    else:
                        return HttpResponse('Unsupported file type')

                    return response
                else:
                    return HttpResponse('No file was uploaded.')
            elif encryption_type == 'decrypt_text':
                print(encryption_type)
                # If no file was uploaded, mask the entered text and display it on the page
                masked_text = unmask_text(text)
                return render(request, 'masker.html', {'form': form, 'masked_text': masked_text, 'text': text})
    else:
        form = MaskerForm()
        return render(request, 'masker.html', {'form': form})

    
def aboutus(request):
    return render(request, 'aboutus.html')